from docx import Document
from datetime import datetime
import pandas as pd

class ContractFiller:
    def __init__(self, template_path, output_path, excel_path):
        self.template_path = template_path
        self.output_path = output_path
        self.excel_path = excel_path
        self.doc = None
        self.current_value_index = 0
        self.placeholder = '_' * 14
        self.date_placeholder = "__ __ ____"
        self.values = []
        
    def load_data(self):
        """Загрузка данных из Excel файла"""
        df = pd.read_excel(self.excel_path)
        self.values = df.iloc[:, 1].astype(str).tolist()
        
    def _copy_font_attributes(self, source_font, target_font):
        """Копирование атрибутов шрифта"""
        attributes = ['name', 'size', 'bold', 'italic', 'underline', 'color']
        for attr in attributes:
            try:
                setattr(target_font, attr, getattr(source_font, attr))
            except AttributeError:
                pass
    
    def _process_paragraph(self, paragraph):
        """Обработка отдельного параграфа"""
        if self.current_value_index >= len(self.values):
            return
            
        original_runs = paragraph.runs
        paragraph.text = ''
        
        for run in original_runs:
            text = run.text
            new_text = []
            
            while self.placeholder in text:
                split_pos = text.find(self.placeholder)
                before = text[:split_pos]
                replacement = text[split_pos:split_pos+14]
                
                if before:
                    new_text.append({'text': before, 'font': run.font})
                
                if replacement == self.placeholder and self.current_value_index < len(self.values):
                    new_text.append({'text': self.values[self.current_value_index], 'font': run.font})
                    self.current_value_index += 1
                else:
                    new_text.append({'text': replacement, 'font': run.font})
                
                text = text[split_pos+14:]
            
            if text:
                new_text.append({'text': text, 'font': run.font})

            for part in new_text:
                new_run = paragraph.add_run(part['text'])
                self._copy_font_attributes(part['font'], new_run.font)
    
    def _replace_date_placeholder(self, paragraph):
        """Замена плейсхолдера даты"""
        if self.date_placeholder in paragraph.text:
            current_date = datetime.today().strftime("%d.%m.%Y")
            paragraph.text = paragraph.text.replace(self.date_placeholder, current_date)
    
    def _process_tables(self):
        """Обработка таблиц в документе"""
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_date_placeholder(paragraph)
                        self._process_paragraph(paragraph)
    
    def generate_document(self):
        """Основной метод генерации документа"""
        self.doc = Document(self.template_path)
        
        # Обработка обычных параграфов
        for paragraph in self.doc.paragraphs:
            self._replace_date_placeholder(paragraph)
            self._process_paragraph(paragraph)
        
        # Обработка таблиц
        self._process_tables()
        
        # Сохранение документа
        self.doc.save(self.output_path)
    
    def run(self):
        """Полный цикл обработки"""
        self.load_data()
        self.generate_document()

if __name__ == "__main__":
    # Конфигурация путей
    config = {
        'template_path': "C:\\Users\\mi\\Documents\\Diplom_Ali4i4\\Auto_Tender_System\\NDA.docx",
        'output_path': "C:\\Users\\mi\\Documents\\Diplom_Ali4i4\\Auto_Tender_System\\готовый_NDA.docx",
        'excel_path': "C:\\Users\\mi\\Documents\\Diplom_Ali4i4\\Auto_Tender_System\\Карточка_контрагента.xlsx"
    }
    
    # Создание и запуск обработчика
    filler = ContractFiller(**config)
    filler.run()